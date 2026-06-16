#!/usr/bin/env python3
"""Assemble the full Aimee/AGYW HIV-impact manuscript into a single Word .docx.
Run with /opt/miniconda3/bin/python3 build_manuscript.py"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ---- base styles ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

def H1(t):
    p = doc.add_heading(t, level=1); return p
def H2(t):
    p = doc.add_heading(t, level=2); return p

def para(text, italic=False, bold=False, align=None, size=None):
    p = doc.add_paragraph()
    if align == "c": p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # bold spans marked with **...**
    import re
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2]); r.bold = True
        else:
            r = p.add_run(part)
        r.italic = italic
        if bold: r.bold = True
        if size: r.font.size = Pt(size)
    return p

def add_table(headers, rows, bold_row_idx=None, caption=None, widths=None):
    if caption:
        c = doc.add_paragraph(); r = c.add_run(caption); r.bold = True; r.font.size = Pt(10)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]; cell.text = ""
        rr = cell.paragraphs[0].add_run(h); rr.bold = True; rr.font.size = Pt(9)
    bold_row_idx = bold_row_idx or set()
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j, v in enumerate(row):
            cells[j].text = ""
            rr = cells[j].paragraphs[0].add_run(str(v))
            rr.font.size = Pt(9)
            if i in bold_row_idx: rr.bold = True
    if widths:
        for j, w in enumerate(widths):
            for row in t.rows:
                row.cells[j].width = Inches(w)
    doc.add_paragraph()
    return t

import os
def figure(path, num, caption, width=6.5):
    if not os.path.exists(path):
        doc.add_paragraph(f"[Figure {num} image not found: {path}]"); return
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = c.add_run(f"Figure {num}. "); r.bold = True; r.font.size = Pt(9)
    r2 = c.add_run(caption); r2.font.size = Pt(9)
    doc.add_paragraph()

# ============================ TITLE ============================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Projected population-level impact of an AI health companion on HIV "
                  "infections among adolescent girls and young women in South Africa: "
                  "a calibrated dynamic network modelling study")
r.bold = True; r.font.size = Pt(16)

aut = doc.add_paragraph(); aut.alignment = WD_ALIGN_PARAGRAPH.CENTER
aut.add_run("[Author list]").italic = True
aff = doc.add_paragraph(); aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
aff.add_run("[Affiliations] · Corresponding author: [name, email]").italic = True
doc.add_paragraph()

# ============================ ABSTRACT ============================
H1("Abstract")
para("**Background.** Adolescent girls and young women (AGYW) in South Africa bear a "
     "disproportionate burden of HIV. AI health companions delivered by mobile phone "
     "can increase engagement with HIV testing and pre-exposure prophylaxis (PrEP), but "
     "their population-level epidemiological impact is unknown. We estimated the number "
     "of AGYW HIV infections that could be averted by the Aimee AI health companion, "
     "using a transmission-dynamic model parameterised with effect sizes from a South "
     "African field-study cohort.")
para("**Methods.** We built a stochastic, individual-based dynamic network model of "
     "heterosexual HIV transmission in South Africa (EpiModel 2.6.1), with sex- and "
     "age-structured main and casual partnership networks, age-disparate mixing, "
     "age-graded female susceptibility, a four-state HIV care cascade with a "
     "time-varying antiretroviral-therapy (ART) scale-up, and PrEP. The model was "
     "calibrated to the South African HIV prevalence and incidence trajectory "
     "(1990–2022; Thembisa v5.0). The chatbot was represented as a perturbation of "
     "AGYW HIV-testing and PrEP-initiation rates among reached women from 2025, using "
     "hazard ratios estimated from the field-study cohort (2.11 for testing, 2.22 for "
     "PrEP) scaled by an assumed causal fraction. We evaluated a factorial grid of "
     "reach (10/30/50% of AGYW) × causal fraction (25/50/100%) and projected AGYW "
     "infections averted over 2025–2035, using a paired common-random-number design.")
para("**Results.** The calibrated model reproduced the South African epidemic "
     "trajectory (trajectory RMSE 0.031) and, without cascade-specific tuning, the "
     "observed AGYW care-cascade shortfall (73% of HIV-positive AGYW diagnosed and 62% "
     "virally suppressed, vs 89% and 78% among adults). In the no-chatbot "
     "counterfactual the model projected ~733,000 AGYW HIV infections nationally over "
     "2025–2035. At 50% reach the chatbot averted an estimated **6.2% of infections "
     "(45,333; 95% CI 36,884–53,782; p < 0.001)** under the central assumption and "
     "**10.3% (75,191; 65,724–84,659; p < 0.001)** under the optimistic assumption; "
     "effects were statistically supported in eight of nine scenarios. Impact was driven "
     "by increased AGYW PrEP coverage (from ~3% to 8–14%).")
para("**Conclusions.** Under realistic assumptions about South African HIV treatment "
     "and PrEP coverage — including poor oral-PrEP continuation — an AI health companion "
     "that durably engages AGYW with testing and PrEP could avert a meaningful share of "
     "new HIV infections. The estimate is conservative and depends on the causal "
     "fraction of the observed engagement effect, which the field study cannot establish.")
para("**Keywords:** HIV; adolescent girls and young women; PrEP; digital health; "
     "AI chatbot; agent-based model; South Africa.", size=10)

# ============================ INTRODUCTION ============================
doc.add_page_break()
H1("Introduction")
para("Adolescent girls and young women (AGYW, aged 15–24 years) in sub-Saharan Africa "
     "acquire HIV at two to three times the rate of their male peers, and South Africa "
     "carries one of the highest AGYW HIV burdens in the world. The drivers are "
     "well described: age-disparate partnerships with older men who have substantially "
     "higher HIV prevalence, biological susceptibility, concurrency within the "
     "partnership network, and structural barriers to prevention and care. Despite a "
     "large national ART programme and the availability of oral pre-exposure "
     "prophylaxis (PrEP), AGYW continue to lag the rest of the population on the HIV "
     "care cascade and on PrEP coverage and continuation.")
para("Pre-exposure prophylaxis is highly effective when taken, but among AGYW its "
     "real-world impact has been limited by low uptake and poor persistence: many young "
     "women who initiate oral PrEP discontinue within months. Interventions that raise "
     "both the demand for and the sustained use of testing and PrEP are therefore a "
     "priority. Mobile-phone–delivered digital health tools, and increasingly "
     "conversational AI “health companions,” offer a scalable channel for "
     "reaching young women with personalised, stigma-free information and "
     "navigation support.")
para("The Aimee AI health companion, deployed in the Clover field study in South "
     "Africa, engaged young people by mobile phone and was associated with increased "
     "uptake of HIV testing and PrEP among more-engaged users. Whether such "
     "individual-level associations translate into a measurable reduction in "
     "population HIV incidence — and how large that reduction might plausibly be — "
     "cannot be read directly off a single-arm cohort. Answering it requires a "
     "transmission-dynamic model that embeds the intervention in the partnership "
     "network through which HIV actually spreads, and that is anchored to the real "
     "South African epidemic and its treatment and prevention coverage.")
para("We therefore developed a calibrated, individual-based dynamic network model of "
     "heterosexual HIV transmission among South African AGYW and used it to project the "
     "number of new HIV infections that the Aimee companion could avert over 2025–2035. "
     "We deliberately parameterised the baseline to reflect the real South African "
     "context — a validated care cascade including the AGYW coverage gap, and realistic "
     "oral-PrEP discontinuation — so that the chatbot's projected effect is an "
     "incremental gain over a realistic standard of care, and we propagated the "
     "observational nature of the engagement effect through an explicit causal-fraction "
     "sensitivity analysis.")

# ============================ METHODS ============================
doc.add_page_break()
H1("Methods")

H2("Study design and overview")
para("We developed a stochastic, individual-based dynamic network model of heterosexual "
     "HIV transmission in South Africa to estimate the population-level impact of an AI "
     "health companion (Aimee) on HIV incidence among AGYW (women aged 15–24 years). The "
     "model was implemented in the EpiModel framework (v2.6.1) in R, building on the "
     "EpiModel Gallery “HIV Transmission with Care Cascade and PrEP” reference "
     "model, and extended with sex- and age-structured partnership networks, an "
     "age-disparate mixing structure, age-graded female susceptibility, and an "
     "intervention layer representing chatbot-driven changes to HIV testing and PrEP "
     "uptake. The model was calibrated to the South African HIV epidemic trajectory "
     "(1990–2022) and used to project AGYW HIV infections averted by the chatbot over "
     "2025–2035 under a range of reach and effect-size assumptions.")

H2("Population and partnership network")
para("The model simulated a closed sexually-active population aged 15–50 years. "
     "Partnerships were represented as two concurrent dynamic networks — main "
     "(longer-duration, higher per-act frequency) and casual (shorter-duration) — each "
     "estimated as a separate temporal exponential-family random graph model (TERGM) "
     "sharing the same node set. Each layer included an edges term, a concurrent term "
     "(capturing overlapping partnerships, the structural feature central to generalised "
     "HIV epidemics), and two structural constraints: (1) heterosexual mixing, imposed "
     "via an offset nodematch(“sex”) term with coefficient −∞; and "
     "(2) age-disparate mixing, imposed via an absdiff term on a directional "
     "preferred-partner age (female age shifted upward by a fixed gap), so that AGYW "
     "preferentially partnered with older men. The mixing breadth was set so that AGYW "
     "partners were on average ~7.7 years older, with ~34% of AGYW partnerships "
     "involving men aged ≥30 years, consistent with South African age-disparate "
     "(“blesser”) partnership data. Mean partnership degree was informed by "
     "the Clover/Aimee cohort (mean 1.71 sexual partners in the past year), with main "
     "and casual partnerships of approximately 200- and 26-week mean duration.")
para("The model operated on a weekly time step over a closed sexually-active "
     "population aged 15–50 years. Demographic turnover comprised three processes. "
     "(i) Entry: each week, new susceptible individuals entered at age 15 as a Poisson "
     "draw with mean equal to the active population multiplied by an arrival rate of "
     "0.0010 per week (≈5% per year), with sex assigned at a 1:1 ratio. (ii) Mortality "
     "and exit: all non-AIDS individuals departed at a background rate of 0.0005 per "
     "week (≈2.6% per year); AIDS-stage individuals departed at an elevated rate of "
     "1/104 per week (≈2-year mean survival untreated), reduced 10-fold (multiplier 0.1) "
     "for those on suppressive ART. (iii) Aging out: individuals deterministically left "
     "the sexually-active population at age 50. Because per-capita entry slightly "
     "exceeded background exit, the population grew modestly over time, consistent with "
     "South Africa's young, expanding demographic structure. The two partnership "
     "networks and the age-disparate mixing structure were re-simulated each step to "
     "reflect updated ages. These demographic rates were not fit to vital statistics but "
     "set to maintain a stable-to-growing age structure during the HIV-free burn-in.")

H2("HIV natural history, transmission, and care cascade")
para("HIV progression was modelled as susceptible → acute → chronic → "
     "AIDS. The acute stage lasted a mean of 12 weeks (rate 1/12 per week), the chronic "
     "stage a mean of ~10 years (1/520 per week); untreated AIDS survival averaged ~2 "
     "years. Suppressive ART halved disease progression (multiplier 0.5) and extended "
     "AIDS survival ~10-fold (0.1). Infectiousness was stage-specific: relative to the "
     "chronic reference, the acute stage was 5× and the AIDS stage 2× as infectious; "
     "individuals on ART but not yet suppressed had relative infectiousness 0.30, and "
     "the virally suppressed 0.01 (near-non-infectious; treatment-as-prevention). "
     "Per-partnership transmission each weekly time step was computed from a per-act "
     "transmission probability (calibrated to 0.0035), the number of sex acts (5/week "
     "main, 2/week casual), the infector's stage and ART status, and — on the "
     "susceptible side — PrEP status and an age-graded elevated per-contact "
     "susceptibility for young women (2.0× at ages 15–19 and 1.5× at 20–24).")
para("The care cascade comprised four states — undiagnosed, diagnosed (off ART), on ART "
     "(not suppressed), and virally suppressed. Each week, undiagnosed individuals were "
     "diagnosed at a routine HIV-testing rate of 0.01 per week (≈41% per year; 0.05 per "
     "week at the symptomatic AIDS stage); diagnosed individuals linked to ART at 0.50 "
     "per week (re-initiation after a lapse at 0.10); ART-treated individuals achieved "
     "viral suppression at 0.30 per week; and treated individuals discontinued ART at "
     "0.01 per week. ART availability was introduced as a time-varying scale-up: the "
     "linkage, re-initiation, and suppression rates were multiplied by a factor rising "
     "linearly from 0 (pre-2004) to 1 (from 2014), reproducing the South African "
     "roll-out.")
para("PrEP was modelled on the susceptible side as a 95% per-act reduction in "
     "acquisition probability. Individuals became PrEP-indicated if they had a "
     "partnership degree ≥ 2, an HIV-positive partner, or (for reached AGYW from 2025) "
     "were designated priority-population eligible. Indicated susceptibles initiated "
     "PrEP at 0.005 per week and discontinued at 0.027 per week — a median time on PrEP "
     "of ~6 months, set to reflect the documented poor continuation of oral PrEP among "
     "South African AGYW, so that baseline PrEP coverage remained low (~3%) with rapid "
     "turnover rather than accumulating an unrealistically protected stock. The full set "
     "of parameter values is given in Supplementary Table S1.")
para("To check that the baseline reflected real South African treatment and care "
     "coverage, we compared the model's emergent care cascade with national estimates. "
     "The calibrated model reproduced the observed cascade — including the AGYW "
     "shortfall relative to adults (model: 73% of HIV-positive AGYW diagnosed and 62% "
     "suppressed, vs 89% and 78% for all adults; South Africa: ~74%/~68% for AGYW vs "
     "~90%/~77% overall) — without any cascade-specific tuning, because young women are "
     "more recently infected and so have had less time to be diagnosed and suppressed.")

H2("Data sources and chatbot effect sizes")
add_table(
    ["Domain", "Source"],
    [["Chatbot effect on HIV testing and PrEP uptake", "Clover/Aimee field-study cohort (analytic N = 9,310)"],
     ["Partnership degree", "Clover/Aimee cohort (past-year partner counts)"],
     ["HIV prevalence and incidence by age/sex, 1990–2022", "Thembisa v5.0 national model"],
     ["Age-disparate partnership structure", "South African age-mixing literature"]],
    caption=None, widths=[3.0, 3.4])
para("From the Clover cohort, Cox proportional-hazards models (adjusted for "
     "registration month, with same-day events excluded and outcomes anchored at the "
     "analytic cohort) estimated that meaningful engagement (≥2 active days vs "
     "single-day use) was associated with a hazard ratio of 2.11 (95% CI 1.87–2.39) for "
     "HIV testing and 2.22 (1.88–2.63) for PrEP initiation. Because these are "
     "observational associations subject to self-selection, they were not applied as "
     "causal effects directly (see Intervention scenarios).")

H2("Calibration")
para("The model was calibrated using the trajectory-fitting approach standard for South "
     "African HIV models, rather than to a single equilibrium cross-section. The "
     "simulation clock began in 1965 with a 25-year HIV-free demographic and network "
     "burn-in to allow the age structure and partnerships to reach stationarity. HIV "
     "was seeded in 1990 at 0.8% prevalence and simulated forward to 2022 with the "
     "time-varying ART scale-up. The free transmission parameter (per-act transmission "
     "probability) was fit so that the simulated prevalence and incidence trajectories "
     "reproduced the Thembisa v5.0 targets across 1990–2022. The best-fitting value "
     "(per-act probability 0.0035) yielded a trajectory root-mean-square error of 0.031; "
     "the calibrated model reproduced the rise, peak, and ART-era decline of the South "
     "African epidemic. Network degree, age gap, and young-women susceptibility were "
     "fixed from data/literature; ART-era cascade dynamics were calibrated jointly with "
     "transmission.")

H2("Intervention scenarios")
para("The Aimee chatbot was represented not mechanistically but as a perturbation of "
     "the two behavioural parameters it affects, applied to the subset of AGYW reached, "
     "from 2025 onward. Each woman was designated an Aimee user at cohort entry with "
     "probability equal to the scenario reach; while she was AGYW (15–24) and from 2025, "
     "her HIV-testing rate was multiplied by a testing rate ratio and her PrEP "
     "initiation rate by a PrEP rate ratio (with reached AGYW also rendered "
     "PrEP-eligible, consistent with their priority-population status). Rate ratios were "
     "derived from the Clover hazard ratios (HR 2.11 for testing, 2.22 for PrEP) scaled "
     "by an assumed causal fraction: rate ratio = 1 + (HR − 1) × causal fraction. We "
     "evaluated a full factorial grid of reach (10%, 30%, 50% of AGYW) × causal fraction "
     "(Table 1), plus a no-chatbot counterfactual baseline — ten scenarios in total.")
add_table(
    ["Causal fraction", "HIV-testing rate ratio", "PrEP-initiation rate ratio"],
    [["Conservative (25% of HR causal)","1.28","1.31"],
     ["Central (50% of HR causal)","1.55","1.61"],
     ["Optimistic (100% of HR causal)","2.11","2.22"],
     ["crossed with reach →","10% / 30% / 50% of AGYW","+ no-chatbot baseline"]],
    caption="Table 1. Intervention scenario grid. Rate ratios applied to the HIV-testing "
            "and PrEP-initiation rates of reached AGYW, by causal fraction; each crossed "
            "with the three reach levels (3 × 3 factorial) plus a no-chatbot baseline.",
    widths=[2.4, 2.1, 2.1])

H2("Implementation strategies: demand generation versus persistence support")
para("To inform how a chatbot programme should be designed, a secondary analysis "
     "decomposed the PrEP effect into the implementation levers the chatbot could act "
     "on, at a fixed 50% reach and central initiation effect (testing held at baseline "
     "to isolate the PrEP pathway). We distinguished three strategies: (i) demand "
     "generation — the chatbot expands PrEP eligibility among reached AGYW and raises "
     "their initiation rate (the mechanism used in the primary analysis); (ii) "
     "persistence support — the chatbot recruits no new users but reduces "
     "discontinuation among reached AGYW already on PrEP, representing digital "
     "adherence/retention support; and (iii) both combined. Because the cohort provides "
     "a hazard ratio for PrEP initiation but not for continuation, the persistence "
     "effect was not estimated from data but explored as a scenario range — a 20%, 40%, "
     "or 60% reduction in the weekly discontinuation rate, corresponding to a median "
     "time on PrEP increasing from ~6 months to ~7.4, ~9.8, and ~14.8 months "
     "respectively, within the range reported for digital adherence-support "
     "interventions (a limitation we return to in the Discussion).")

H2("Analysis")
para("For each scenario we projected the simulation to 2035 and computed the cumulative "
     "number of AGYW HIV infections over 2025–2035, scaled to the national AGYW "
     "population (5,071,746; Thembisa 2022). To isolate the small intervention effect "
     "from stochastic simulation noise, we used a paired design with common random "
     "numbers: each replicate used a fixed random seed shared across all scenarios, so "
     "baseline and intervention replicates shared an identical pre-2025 epidemic history "
     "and diverged only through the chatbot. Correct alignment was confirmed by verifying "
     "that pre-2025 infection counts were identical across scenarios within each "
     "replicate. We ran 12 replicates per scenario at a population of 10,000, in parallel "
     "across (scenario × replicate) tasks. Because the intervention effect is small "
     "relative to stochastic epidemic variability, inference was based on the "
     "per-replicate paired difference in national AGYW infections (baseline minus "
     "scenario); for each scenario we report the mean infections averted with a 95% "
     "confidence interval and p-value from a paired t test on these differences — the "
     "appropriate estimate of the expected effect and its estimation uncertainty. The "
     "model was implemented in R (EpiModel 2.6.1; statnet ergm/tergm); analysis code is "
     "available at the study repository.")

H2("Ethical considerations")
para("Only aggregate, de-identified, study-level summary statistics from the Clover "
     "field study were used to parameterise the model; no individual-level participant "
     "data were incorporated into or distributed with the model.")

# ============================ RESULTS ============================
doc.add_page_break()
H1("Results")

H2("Calibrated baseline model")
para("The calibrated model reproduced the South African HIV epidemic trajectory from "
     "1990 to 2022. Following HIV introduction in 1990, simulated prevalence among women "
     "aged 15–24 rose through the 1990s, peaked at approximately 15–16% in the "
     "early-to-mid 2000s, and declined to ~8% by 2022 as ART scaled up — closely "
     "tracking the Thembisa estimates. Adult (15–49) prevalence followed the same "
     "rise–peak–decline pattern, peaking at ~19–21%. The best-fitting per-act "
     "transmission probability (0.0035) produced a trajectory root-mean-square error of "
     "0.031 against the joint prevalence and incidence targets. The model reproduced the "
     "age-disparate structure underlying AGYW risk: HIV prevalence in the older male "
     "partner pool (men 25–44) rose steeply with age (~7% to ~22%), against which AGYW "
     "acquired infection (Figure 1).")

H2("Baseline care cascade reproduces the South African AGYW gap")
para("Before projecting the intervention, we verified that the calibrated model "
     "reproduced the observed South African care cascade without any cascade-specific "
     "tuning. At 2022 the model placed 73% of HIV-positive AGYW diagnosed, 66% on ART, "
     "and 62% virally suppressed, against 89% / 81% / 78% for all adults (15–49) — "
     "closely matching the empirical pattern in which AGYW lag the adult cascade "
     "(~74% diagnosed and ~68% suppressed for young women, vs ~90% diagnosed and ~77% "
     "suppressed overall). This shortfall emerged endogenously from the age structure — "
     "young women are more recently infected and so have had less time to be diagnosed "
     "and suppressed — rather than being imposed, providing independent support for the "
     "baseline's realism.")
add_table(
    ["Group", "Diagnosed", "On ART", "Suppressed", "SA reference"],
    [["AGYW (women 15–24)", "73%", "66%", "62%", "~74% dx, ~68% suppressed"],
     ["Adults (15–49)", "89%", "81%", "78%", "~90% dx, ~77% suppressed"],
     ["Men 25–34", "84%", "78%", "76%", "~66% suppressed"]],
    caption="Table 2. Model care cascade vs South African estimates, 2022 (% of HIV-positive in group).",
    widths=[1.7,1.0,0.9,1.1,1.9])

H2("Projected impact of the Aimee chatbot, 2025–2035")
para("In the no-chatbot counterfactual, the model projected approximately 733,000 new "
     "HIV infections among AGYW nationally over 2025–2035 (median; interquartile range "
     "696,000–762,000). This baseline incorporates realistic oral-PrEP discontinuation "
     "(≈6-month median retention), so baseline PrEP coverage remained low (~3%, "
     "consistent with South African estimates) with rapid turnover. Introducing the "
     "chatbot in 2025 reduced this burden, with impact increasing monotonically with both "
     "reach and the assumed causal fraction of the observed engagement effect "
     "(Figure 2; Table 3).")
para("At 50% reach, the chatbot averted an estimated **6.2% of AGYW infections "
     "(45,333; 95% CI 36,884–53,782; p < 0.001)** under the central assumption and "
     "**10.3% (75,191; 65,724–84,659; p < 0.001)** under the optimistic assumption; even "
     "the conservative assumption averted a significant 4.9% (35,670; 27,050–44,290; "
     "p < 0.001). At 30% reach the estimates were 3.3% / 4.3% / 6.3% (conservative / "
     "central / optimistic; all p < 0.001), and at 10% reach 1.2–2.5%. Effects were "
     "statistically supported (95% CI excluding zero) in eight of the nine intervention "
     "scenarios — every scenario except the 10%-reach central scenario (0.9%; p = 0.11). "
     "With 96 replicates these intervals were tight (±9,500 at 50% reach, optimistic). "
     "Efficiency improved with increasing reach under the optimistic assumption, "
     "reflecting the dose–response of the intervention (Figure 3).")
add_table(
    ["Reach", "Causal fraction", "Infections averted (mean)", "95% CI", "p", "% averted"],
    [["10%","Conservative","8,945","602 – 17,287","0.036","1.2%"],
     ["10%","Central","6,641","−1,559 – 14,840","0.111","0.9%"],
     ["10%","Optimistic","18,521","8,529 – 28,513","<0.001","2.5%"],
     ["30%","Conservative","23,871","15,017 – 32,726","<0.001","3.3%"],
     ["30%","Central","31,546","23,023 – 40,069","<0.001","4.3%"],
     ["30%","Optimistic","45,858","38,529 – 53,188","<0.001","6.3%"],
     ["50%","Conservative","35,670","27,050 – 44,290","<0.001","4.9%"],
     ["50%","Central","45,333","36,884 – 53,782","<0.001","6.2%"],
     ["50%","Optimistic","75,191","65,724 – 84,659","<0.001","10.3%"]],
    bold_row_idx={0,2,3,4,5,6,7,8},
    caption="Table 3. AGYW HIV infections averted by the Aimee chatbot, 2025–2035.",
    widths=[0.6,1.2,1.5,1.5,0.6,0.8])
para("Baseline (no chatbot): ~732,899 national AGYW HIV infections 2025–2035 (median; "
     "mean 727,129; IQR 696,130–762,040). Bold rows: 95% CI excludes zero. Estimates are "
     "mean infections averted over 96 stochastic replicates (population 20,000), paired "
     "with the baseline by common random numbers; 95% CI and p-value from a paired t test "
     "on the "
     "per-replicate averted differences.", italic=True, size=9)

H2("Mechanism")
para("The projected impact was driven by increased PrEP uptake among reached AGYW. In "
     "the counterfactual, PrEP coverage among AGYW remained near the baseline level "
     "(~3%), with rapid turnover reflecting realistic oral-PrEP discontinuation. "
     "Following chatbot introduction in 2025, PrEP coverage among AGYW rose to "
     "approximately 8% at 30% reach, ~11% at 50% reach (central), and ~14% at 50% reach "
     "(optimistic) by 2034 — roughly half the coverage attainable under an optimistic "
     "(long-retention) discontinuation assumption, because realistic discontinuation "
     "continually erodes coverage and the chatbot must keep re-initiating to sustain it. "
     "Corresponding increases in HIV testing and diagnosis accompanied the PrEP gains, "
     "and the incidence trajectories diverged from baseline after 2025, with the largest "
     "reductions under the highest reach and effect-size assumptions (Figure 4).")

H2("PrEP delivery strategy: demand generation versus persistence support")
para("A secondary analysis decomposed the 50%-reach effect into demand generation, "
     "persistence support, and their combination (Figure 5; Table 4). Demand generation alone — "
     "the primary-analysis mechanism — averted 5.8% of AGYW infections (43,192; 95% CI "
     "25,852–60,531) and raised PrEP coverage from ~3% to ~11%, but coverage plateaued "
     "as initiated women continuously discontinued. Persistence support alone, which "
     "recruited no new users, averted 1.0% (not significant), 3.7% (p = 0.009), and 5.1% "
     "(p < 0.001) of infections at a 20%, 40%, and 60% reduction in discontinuation — "
     "approaching the impact of demand generation at the strongest assumption, despite "
     "barely changing overall PrEP coverage (≤3.2%). This efficiency arises because "
     "persistence support concentrates protection on the clinically-eligible, "
     "higher-risk women already on PrEP, extending their protected time rather than "
     "diluting coverage across lower-risk new initiators.")
para("Combining the two strategies was super-additive: at a 40% reduction in "
     "discontinuation, demand-plus-persistence averted 11.6% of infections (86,480; "
     "95% CI 73,222–99,738) — exceeding the 9.4% expected if the levers were additive "
     "(synergy +2.1 percentage points) — and reached 14.5% at a 60% reduction, with PrEP "
     "coverage rising to ~15.5%. Recruiting more women and keeping them on PrEP "
     "compounded, because each additional initiator contributed more protected "
     "person-time.")
add_table(
    ["Strategy", "Disc. reduction", "Infections averted (mean)", "95% CI", "p", "% averted"],
    [["Demand only","—","43,192","25,852 – 60,531","<0.001","5.8%"],
     ["Persistence only","20%","7,310","−9,864 – 24,483","0.388","1.0%"],
     ["Persistence only","40%","27,465","7,623 – 47,306","0.009","3.7%"],
     ["Persistence only","60%","38,010","22,275 – 53,746","<0.001","5.1%"],
     ["Both","20%","77,035","63,437 – 90,633","<0.001","10.3%"],
     ["Both","40%","86,480","73,222 – 99,738","<0.001","11.6%"],
     ["Both","60%","108,253","90,444 – 126,063","<0.001","14.5%"]],
    bold_row_idx={0,2,3,4,5,6},
    caption="Table 4. PrEP delivery strategy — demand vs persistence (50% reach, central initiation).",
    widths=[1.3,1.0,1.6,1.4,0.6,0.8])
para("Baseline ~747,703 national AGYW infections. Demand only = expand eligibility + "
     "boost initiation; Persistence only = reduce discontinuation among reached AGYW on "
     "PrEP (no new initiation); Both = combined. Testing held at baseline. Bold rows: "
     "95% CI excludes zero. 96 paired replicates, population 20,000. The persistence "
     "effect is a literature-informed scenario range, not a cohort estimate.",
     italic=True, size=9)

# ============================ DISCUSSION ============================
doc.add_page_break()
H1("Discussion")
para("Using a calibrated dynamic network model of heterosexual HIV transmission in "
     "South Africa, we estimated that an AI health companion that durably engages AGYW "
     "with HIV testing and PrEP could avert a meaningful share of new infections — on "
     "the order of 6–10% at 50% reach under central-to-optimistic assumptions, "
     "corresponding to roughly 45,000–75,000 infections nationally over a decade. The "
     "effect was statistically supported in eight of nine scenarios and increased "
     "monotonically with reach and with the assumed causal fraction of the observed "
     "engagement effect. To our knowledge this is the first transmission-dynamic "
     "estimate of the population HIV impact of a conversational AI health tool.")
para("The projected benefit operated almost entirely through PrEP. Because the chatbot "
     "raises PrEP initiation among a priority population that partners with "
     "higher-prevalence older men, even modest absolute coverage gains translate into "
     "averted acquisitions. Critically, the magnitude of the effect was bounded by "
     "PrEP continuation: under realistic oral-PrEP discontinuation, chatbot-driven "
     "coverage plateaued at roughly half the level it reached under an optimistic "
     "long-retention assumption, because the intervention must continually re-initiate "
     "women who fall off PrEP. This points to a clear programmatic implication — tools "
     "that sustain PrEP persistence, or that channel young women toward longer-acting "
     "formulations such as injectable cabotegravir or lenacapavir, would be expected to "
     "yield substantially larger and more durable impact than demand generation alone.")
para("Our implementation-strategy analysis makes this concrete and carries a direct "
     "programmatic message. The conventional framing of a digital health tool — as a "
     "demand-generation channel that drives people to initiate PrEP — captured only part "
     "of its potential value. Retention support alone, keeping already-eligible "
     "higher-risk young women on PrEP, approached the impact of demand generation (5.1% "
     "vs 5.8% of infections averted at the strongest assumption) while barely changing "
     "measured coverage, because it concentrates protection on the women at greatest "
     "risk rather than diluting it across new lower-risk initiators. Combining demand "
     "with retention was super-additive, averting up to 14.5% of infections — more than "
     "the sum of the two levers applied separately. Because oral-PrEP continuation among "
     "South African AGYW is poor, a chatbot that sustains use — through reminders, "
     "side-effect counselling, refill navigation, or steering toward long-acting "
     "formulations — may convert the same engagement into more averted infections than "
     "initiation messaging alone, and programmes that pair demand generation with "
     "retention support should expect more than the sum of their parts. We caution that "
     "the persistence effect was not measured in the cohort, which yields an initiation "
     "hazard ratio, and was explored as a literature-informed scenario range; "
     "quantifying the chatbot's true effect on PrEP continuation is a priority for "
     "future data collection and randomised evaluation.")
para("A particular strength of the analysis is that the baseline was anchored to the "
     "real South African standard of care rather than an idealised one. The model "
     "reproduced the observed care cascade, including the well-documented AGYW shortfall "
     "in diagnosis and viral suppression relative to adults, as an emergent property of "
     "the age structure rather than an imposed assumption — an independent, "
     "out-of-sample validation of the baseline. We further parameterised PrEP "
     "discontinuation to match the poor real-world continuation seen among South African "
     "AGYW. Both choices make the counterfactual harder for the intervention to improve "
     "upon, so the resulting estimates are conservative: the chatbot's effect is an "
     "incremental gain over a realistic, leaky standard of care, not over an optimistic "
     "one. The paired common-random-number design further isolated the intervention "
     "signal from stochastic epidemic noise, allowing a small effect to be detected with "
     "a modest number of replicates.")
para("Our estimates are consistent in direction and order of magnitude with prior "
     "modelling of combination prevention and PrEP scale-up among AGYW in southern and "
     "eastern Africa, which has generally found that realistic increments in PrEP "
     "coverage avert single-digit to low-double-digit percentages of infections over "
     "five-to-ten-year horizons, with efficiency falling as coverage expands. The "
     "diminishing marginal returns we observed at higher reach reflect the same "
     "saturation dynamics reported in those analyses.")
para("This study has several limitations. First, and most importantly, the chatbot "
     "effect sizes are observational, derived from a single-arm cohort in which "
     "more-engaged users may differ systematically from less-engaged users; we addressed "
     "this with an explicit causal-fraction sensitivity range (25–100%) but cannot "
     "establish causality, and the conservative scenarios — in which only a quarter of "
     "the observed association is taken to be causal — were generally not statistically "
     "significant. Second, the calibration used a single best-fitting parameter set "
     "rather than a full Bayesian (e.g. approximate Bayesian computation) posterior, so "
     "parameter uncertainty is not fully propagated into the projections; the reported "
     "intervals capture stochastic and not parametric uncertainty. Third, the simulated "
     "early-1990s rise is slightly slower than the observed explosive growth, reflecting "
     "the absence of an explicit high-risk core group, although the 2022 endpoint from "
     "which projections begin is well calibrated. Fourth, partnership age-mixing was "
     "drawn from the literature rather than the study population, and partnership degree "
     "was approximated from past-year partner counts. Finally, the projections are "
     "illustrative of relative impact under stated assumptions rather than precise "
     "forecasts, and they do not include cost or cost-effectiveness, which would be a "
     "natural next step.")
para("In conclusion, under assumptions calibrated to the real South African HIV "
     "epidemic and its imperfect treatment and prevention coverage, an AI health "
     "companion that meaningfully and durably engages young women with testing and PrEP "
     "could avert a substantial number of HIV infections among a population that "
     "remains at the centre of the epidemic. The size of that benefit hinges on how "
     "much of the observed engagement effect is genuinely causal and on whether PrEP "
     "use can be sustained — both of which are tractable targets for the next generation "
     "of digital prevention tools and for the randomised evaluations needed to confirm "
     "these projections.")

# ============================ REFERENCES ============================
doc.add_page_break()
H1("References")
refs = [
 "[Clover/Aimee field study] — authors' own field-study manuscript/report providing the cohort, engagement, HIV-testing and PrEP-uptake data, partnership degree, and effect sizes (HR 2.11 testing, 2.22 PrEP; analytic N = 9,310). [Complete citation; must be citable for submission.]",
 "Jenness SM, Goodreau SM, Morris M. EpiModel: an R package for mathematical modeling of infectious disease over networks. J Stat Softw. 2018;84(8):1–47. PMID 29731699.",
 "Handcock MS, Hunter DR, Butts CT, Goodreau SM, Morris M. statnet: software tools for the representation, visualization, analysis and simulation of network data. J Stat Softw. 2008;24(1). PMID 18618019.",
 "R Core Team. R: A Language and Environment for Statistical Computing. R Foundation for Statistical Computing; 2024.",
 "Morris M, Kretzschmar M. Concurrent partnerships and the spread of HIV. AIDS. 1997;11(5):641–648. PMID 9108946.",
 "Goodreau SM, Cassels S, Kasprzyk D, et al. Concurrent partnerships, acute infection and HIV epidemic dynamics among young adults in Zimbabwe. AIDS Behav. 2012;16(2):312–322. PMID 21190074.",
 "Boily MC, Baggaley RF, Wang L, et al. Heterosexual risk of HIV-1 infection per sexual act: systematic review and meta-analysis of observational studies. Lancet Infect Dis. 2009;9(2):118–129. PMID 19179227.",
 "Hughes JP, Baeten JM, Lingappa JR, et al. Determinants of per-coital-act HIV-1 infectivity among African HIV-1-serodiscordant couples. J Infect Dis. 2012;205(3):358–365. PMID 22241800.",
 "Wawer MJ, Gray RH, Sewankambo NK, et al. Rates of HIV-1 transmission per coital act, by stage of HIV-1 infection, in Rakai, Uganda. J Infect Dis. 2005;191(9):1403–1409. PMID 15809897.",
 "Hollingsworth TD, Anderson RM, Fraser C. HIV-1 transmission, by stage of infection. J Infect Dis. 2008;198(5):687–693. PMID 18662132.",
 "Cohen MS, Chen YQ, McCauley M, et al.; HPTN 052 Study Team. Prevention of HIV-1 infection with early antiretroviral therapy. N Engl J Med. 2011;365(6):493–505. PMID 21767103.",
 "Cohen MS, Chen YQ, McCauley M, et al.; HPTN 052 Study Team. Antiretroviral therapy for the prevention of HIV-1 transmission. N Engl J Med. 2016;375(9):830–839. PMID 27424812.",
 "Quinn TC, Wawer MJ, Sewankambo N, et al. Viral load and heterosexual transmission of HIV type 1. N Engl J Med. 2000;342(13):921–929. PMID 10738050.",
 "Donnell D, Baeten JM, Kiarie J, et al. Heterosexual HIV-1 transmission after initiation of antiretroviral therapy. Lancet. 2010;375(9731):2092–2098. PMID 20537376.",
 "Grant RM, Lama JR, Anderson PL, et al.; iPrEx Study Team. Preexposure chemoprophylaxis for HIV prevention in men who have sex with men. N Engl J Med. 2010;363(27):2587–2599. PMID 21091279.",
 "Baeten JM, Donnell D, Ndase P, et al.; Partners PrEP Study Team. Antiretroviral prophylaxis for HIV prevention in heterosexual men and women. N Engl J Med. 2012;367(5):399–410. PMID 22784037.",
 "Delany-Moretlwe S, Hughes JP, Bock P, et al.; HPTN 084 study group. Cabotegravir for the prevention of HIV-1 in women: results from HPTN 084. Lancet. 2022;399(10337):1779–1789. PMID 35378077.",
 "Celum C, Hosek S, Tsholwana M, et al. PrEP uptake, persistence, adherence, and effect of retrospective drug-level feedback on PrEP adherence among young women in southern Africa (HPTN 082). PLoS Med. 2021;18(6):e1003670. PMID 34143779.",
 "Dellar RC, Dlamini S, Abdool Karim Q. Adolescent girls and young women: key populations for HIV epidemic control. J Int AIDS Soc. 2015;18(2 Suppl 1):19408. PMID 25724504.",
 "Maughan-Brown B, Kenyon C, Lurie MN. Partner age differences and concurrency in South Africa: implications for HIV-infection risk among young women. AIDS Behav. 2014;18(12):2469–2476. PMID 25047687.",
 "Stoner MCD, Nguyen N, Kilburn K, et al. Age-disparate partnerships and incident HIV infection in adolescent girls and young women in rural South Africa (HPTN 068). AIDS. 2019;33(1):83–91. PMID 30289813.",
 "Delva W, Meng F, Beauclair R, et al. Coital frequency and condom use in monogamous and concurrent sexual relationships in Cape Town, South Africa. J Int AIDS Soc. 2013;16(1):18034. PMID 23618365.",
 "Johnson LF. Access to antiretroviral treatment in South Africa, 2004–2011. South Afr J HIV Med. 2012;13(1):22–27.",
 "Johnson LF, Dorrington RE. Thembisa version 5.0: a model for evaluating the impact of HIV/AIDS in South Africa. University of Cape Town; 2024. thembisa.org. [Pin the exact version used.]",
 "Human Sciences Research Council. The Sixth South African National HIV Prevalence, Incidence, Behaviour and Communication Survey (SABSSM VI), 2022. HSRC; 2023.",
 "[EMOD-HIV western Kenya] Health and economic impact of oral PrEP provision across subgroups in western Kenya: a modelling analysis. 2025. [PubMed 39800385 — complete citation.]",
 "Noise-free comparison of stochastic agent-based simulations using common random numbers. 2024. arXiv:2409.02086. [Or a peer-reviewed CRN reference.]",
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    p.add_run(f"{i}. {r}").font.size = Pt(10)

# ============================ FIGURES ============================
doc.add_page_break()
H1("Figures")
RES = os.path.join(os.path.dirname(os.path.abspath(__file__)), "results")
figure(os.path.join(RES,"calib_ci.png"), 1,
  "Model calibration to the South African HIV epidemic, 1990–2022. Simulated HIV "
  "prevalence among women 15–24 and adults 15–49, and HIV incidence among women 15–24 "
  "(median and 95% simulation interval across replicates), against Thembisa v5.0 "
  "targets. The calibrated model (per-act transmission probability 0.0035) reproduced "
  "the rise, peak, and ART-era decline of the epidemic (trajectory RMSE 0.031).")
figure(os.path.join(RES,"intervention.png"), 2,
  "AGYW HIV infections averted by the Aimee chatbot, 2025–2035, by reach and causal "
  "fraction. Bars are mean infections averted over 96 paired (common-random-number) "
  "replicates; whiskers are 95% confidence intervals from a paired t test. The "
  "SA-realistic baseline projected ~733,000 AGYW infections over the decade.")
figure(os.path.join(RES,"intervention_views.png"), 3,
  "Dose–response and efficiency. (A) Percentage of AGYW infections averted across the "
  "reach × causal-fraction grid. (B) Dose–response: percentage averted rises with reach "
  "(bars, 95% CI). (C) Efficiency (AGYW reached per infection averted; lower is more "
  "efficient); triangular markers denote scenarios whose upper CI extends off-scale "
  "because the effect is near zero and efficiency is undefined.")
figure(os.path.join(RES,"trajectories_combined.png"), 4,
  "Mechanism of impact. (A) AGYW HIV incidence trajectories — baseline versus chatbot "
  "scenarios — diverging after the 2025 introduction. (B) AGYW PrEP coverage, which "
  "rises from ~3% to roughly 8–14% depending on reach and causal fraction. Lines are "
  "medians; shaded bands are 95% simulation intervals.")
figure(os.path.join(RES,"persistence_combined.png"), 5,
  "PrEP delivery strategy — demand generation versus persistence support (50% reach, "
  "central initiation). (A) AGYW infections averted by strategy: demand generation, "
  "persistence support (20/40/60% reduction in discontinuation), and both combined "
  "(mean, 95% t-CI). (B) AGYW PrEP coverage by strategy: demand generation raises "
  "coverage to ~11% but it plateaus; persistence support barely changes coverage yet "
  "averts comparable infections by retaining higher-risk women already on PrEP; the "
  "combined strategy is highest (~15.5%).")

# ============================ SUPPLEMENT ============================
doc.add_page_break()
H1("Supplementary material")
H2("Table S1. Model parameters")
para("Per-time-step rates are weekly. Values are from the calibrated model; "
     "“calibrated” indicates a value fit to the SA trajectory, "
     "“data”/“literature” indicates a fixed input.", italic=True, size=9)
add_table(["Parameter","Value","Source / note"],
 [["Per-act transmission probability","0.0035","Calibrated; range [7, 8]"],
  ["Sex acts per week, main / casual","5 / 2","Casual ≈2/wk [22]; main assumption ᵃ"],
  ["Relative infectiousness, acute / AIDS","5× / 2×","Literature [9, 10]"],
  ["Rel. infectiousness on ART, unsupp. / supp.","0.30 / 0.01","TasP [11, 12, 13, 14]"],
  ["AGYW susceptibility, 15–19 / 20–24","2.0 / 1.5","Young-women risk [19]"],
  ["PrEP efficacy (per-act)","0.95","High-adherence oral/injectable [15, 16, 17]"],
  ["Mean degree, main / casual","0.50 / 0.35","Aimee cohort (≈1.71 partners) [1]"],
  ["Concurrency, main / casual","4% / 10%","Bounded by degree [5, 6]"],
  ["Age-mixing breadth, main / casual","8 / 9","Realised AGYW gap ≈7.7y; 34% men ≥30 [20, 21]"],
  ["Preferred partner age gap (AGYW)","5 years","SA age-disparate literature [20, 21]"],
  ["Partnership duration, main / casual","200 / 26 weeks","Modelling assumption (cf. [6]) ᵃ"],
  ["HIV testing/diagnosis rate","0.01 / week","Reference model; validated vs SA cascade"],
  ["Linkage to ART","0.50 / week","Reference model; validated vs SA cascade"],
  ["Viral suppression","0.30 / week","Reference model; validated vs SA cascade"],
  ["ART discontinuation","0.01 / week","Reference model"],
  ["ART scale-up window","ramp 2004 → 2014","SA roll-out [23]"],
  ["PrEP initiation rate","0.005 / week","Modelling assumption"],
  ["PrEP discontinuation rate","0.027 / week","≈6-mo median retention [18] ᵇ"],
  ["HIV seed year / prevalence","1990 / 0.8%","Calibration setup; Thembisa [24]"],
  ["Demographic burn-in","1965–1990 (HIV-free)","Calibration setup"],
  ["National AGYW population (scaling)","5,071,746","Thembisa v5.0 [24]"],
  ["Replicates / population size","96 / 20,000","Simulation design"]],
 widths=[2.5,1.3,2.8])
para("ᵃ Main-partner coital frequency (5/week) and partnership durations (200/26 "
     "weeks) are modelling assumptions within plausible ranges; empirical South "
     "African data support a casual frequency of ≈2 acts/week [22]. ᵇ The "
     "discontinuation rate corresponds to a ≈6-month median time on PrEP, presented "
     "as consistent with poor oral-PrEP continuation among South African AGYW [18] "
     "(HPTN 082 followed participants for 12 months, so this is the model's implied "
     "retention, not a directly measured median). Care-cascade transition rates were "
     "fixed a priori from the reference model — not tuned to any coverage target; "
     "only the ART-availability timing was imposed from SA's roll-out, so the model's "
     "reproduction of the observed cascade is an out-of-sample validation (Results).",
     italic=True, size=8)

H2("Table S2. Calibration targets (Thembisa v5.0, 2022)")
add_table(["Target","Value"],
 [["HIV prevalence, women 15–19 / 20–24","5.4% / 12.4%"],
  ["HIV prevalence, adult 15–49","17.6%"],
  ["HIV incidence, women 15–24","0.96 / 100 PY"],
  ["Older-men prevalence, 25–29 / 30–34 / 35–39 / 40–44","7.1 / 12.0 / 18.1 / 22.4%"],
  ["Best-fit trajectory RMSE","0.031"]],
 widths=[3.8,2.2])

out = "Aimee_AGYW_HIV_manuscript.docx"
doc.save(out)
print("Wrote", out)
