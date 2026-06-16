#!/usr/bin/env python3
"""Standalone supplementary document: model parameters + calibration targets,
with a SELF-CONTAINED reference list (numbered 1..n, independent of the main paper).
Run: /opt/miniconda3/bin/python3 build_supplement.py"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
n = doc.styles["Normal"]; n.font.name = "Calibri"; n.font.size = Pt(11)
n.paragraph_format.space_after = Pt(6); n.paragraph_format.line_spacing = 1.15

def H1(t): return doc.add_heading(t, level=1)
def H2(t): return doc.add_heading(t, level=2)
def para(text, italic=False, size=None):
    p = doc.add_paragraph()
    import re
    for part in re.split(r"(\*\*.*?\*\*)", text):
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2]); r.bold = True
        else:
            r = p.add_run(part)
        r.italic = italic
        if size: r.font.size = Pt(size)
    return p
def add_table(headers, rows, widths=None, caption=None):
    if caption:
        c = doc.add_paragraph(); r = c.add_run(caption); r.bold = True; r.font.size = Pt(10)
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]; cell.text = ""
        rr = cell.paragraphs[0].add_run(h); rr.bold = True; rr.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for j, v in enumerate(row):
            cells[j].text = ""; cells[j].paragraphs[0].add_run(str(v)).font.size = Pt(9)
    if widths:
        for j, w in enumerate(widths):
            for row in t.rows: row.cells[j].width = Inches(w)
    doc.add_paragraph(); return t

# ---------------- Title ----------------
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Supplementary Material\nProjected population-level impact of an AI health companion "
                  "on HIV infections among adolescent girls and young women in South Africa")
r.bold = True; r.font.size = Pt(14)
para("This supplement lists the full model parameter set, the calibration targets, and a "
     "self-contained reference list. Reference numbers below are specific to this supplement.",
     italic=True, size=10)

# ---------------- Table S1 ----------------
H2("Table S1. Model parameters")
para("Per-time-step rates are weekly. “Calibrated” = fit to the South African "
     "trajectory; otherwise values are fixed inputs. Citations refer to the supplementary "
     "reference list below.", italic=True, size=9)
add_table(["Parameter", "Value", "Source"],
 [["Per-act transmission probability", "0.0035", "Calibrated; range [1, 2]"],
  ["Sex acts per week, main / casual", "5 / 2", "Casual ≈2/wk [3]; main assumption ᵃ"],
  ["Relative infectiousness, acute / AIDS", "5× / 2×", "Wawer 2005; Hollingsworth 2008 [4, 5]"],
  ["Rel. infectiousness on ART, unsupp. / supp.", "0.30 / 0.01", "HPTN 052; Quinn 2000; Donnell 2010 [6–9]"],
  ["AGYW susceptibility, 15–19 / 20–24", "2.0 / 1.5", "Dellar 2015 [10]"],
  ["PrEP efficacy (per-act)", "0.95", "iPrEx; Partners PrEP; HPTN 084 [11–13]"],
  ["Mean degree, main / casual", "0.50 / 0.35", "Aimee cohort (≈1.71 partners) [14]"],
  ["Concurrency, main / casual", "4% / 10%", "Morris & Kretzschmar 1997; Goodreau 2012 [15, 16]"],
  ["Age-mixing breadth, main / casual", "8 / 9", "Realised AGYW gap ≈7.7y; 34% men ≥30 [17, 18]"],
  ["Preferred partner age gap (AGYW)", "5 years", "Maughan-Brown 2014; Stoner 2019 [17, 18]"],
  ["Partnership duration, main / casual", "200 / 26 weeks", "Modelling assumption (cf. [16]) ᵃ"],
  ["HIV testing/diagnosis rate", "0.01 / week", "Reference model; validated vs SA cascade"],
  ["Linkage to ART", "0.50 / week", "Reference model; validated vs SA cascade"],
  ["Viral suppression", "0.30 / week", "Reference model; validated vs SA cascade"],
  ["ART discontinuation", "0.01 / week", "Reference model"],
  ["ART scale-up window", "ramp 2004 → 2014", "Johnson 2012 [19]"],
  ["PrEP initiation rate", "0.005 / week", "Modelling assumption"],
  ["PrEP discontinuation rate", "0.027 / week", "≈6-mo median retention; HPTN 082 [20] ᵇ"],
  ["HIV seed year / prevalence", "1990 / 0.8%", "Calibration setup; Thembisa [21]"],
  ["Demographic burn-in", "1965–1990 (HIV-free)", "Calibration setup"],
  ["National AGYW population (scaling)", "5,071,746", "Thembisa v5.0 [21]"],
  ["Replicates / population size", "96 / 20,000", "Simulation design"]],
 widths=[2.5, 1.3, 2.8])
para("ᵃ Main-partner coital frequency (5/week) and partnership durations (200/26 weeks) "
     "are modelling assumptions within plausible ranges; South African data support a casual "
     "frequency of ≈2 acts/week [3]. ᵇ The 0.027/week rate corresponds to a ≈6-month "
     "median time on PrEP, presented as consistent with poor oral-PrEP continuation among South "
     "African AGYW [20]; HPTN 082 followed participants for 12 months, so this is the model’s "
     "implied retention rather than a directly measured median. Care-cascade transition rates were "
     "fixed a priori (not tuned to any coverage target); only the ART-availability timing was "
     "imposed from South Africa’s roll-out [19], so the model’s reproduction of the observed "
     "cascade is an out-of-sample validation.", italic=True, size=8)

# ---------------- Table S2 ----------------
H2("Table S2. Calibration targets (Thembisa v5.0, 2022)")
add_table(["Target", "Value"],
 [["HIV prevalence, women 15–19 / 20–24", "5.4% / 12.4%"],
  ["HIV prevalence, adult 15–49", "17.6%"],
  ["HIV incidence, women 15–24", "0.96 / 100 PY"],
  ["Older-men prevalence, 25–29 / 30–34 / 35–39 / 40–44", "7.1 / 12.0 / 18.1 / 22.4%"],
  ["Best-fit prevalence-trajectory RMSE", "0.031 (≈3 percentage points)"]],
 widths=[3.8, 2.2])
para("Targets are from the Thembisa v5.0 national model [21]; cross-checked against SABSSM VI "
     "(2022) [22]. The transmission probability was fit to the two prevalence trajectories; "
     "incidence was assessed separately (not combined into the RMSE).", italic=True, size=9)

# ---------------- Supplementary references (self-contained) ----------------
H2("Supplementary references")
refs = [
 "Boily MC, Baggaley RF, Wang L, et al. Heterosexual risk of HIV-1 infection per sexual act: systematic review and meta-analysis of observational studies. Lancet Infect Dis. 2009;9(2):118–129. PMID 19179227.",
 "Hughes JP, Baeten JM, Lingappa JR, et al. Determinants of per-coital-act HIV-1 infectivity among African HIV-1-serodiscordant couples. J Infect Dis. 2012;205(3):358–365. PMID 22241800.",
 "Delva W, Meng F, Beauclair R, et al. Coital frequency and condom use in monogamous and concurrent sexual relationships in Cape Town, South Africa. J Int AIDS Soc. 2013;16(1):18034. PMID 23618365.",
 "Wawer MJ, Gray RH, Sewankambo NK, et al. Rates of HIV-1 transmission per coital act, by stage of HIV-1 infection, in Rakai, Uganda. J Infect Dis. 2005;191(9):1403–1409. PMID 15809897.",
 "Hollingsworth TD, Anderson RM, Fraser C. HIV-1 transmission, by stage of infection. J Infect Dis. 2008;198(5):687–693. PMID 18662132.",
 "Cohen MS, Chen YQ, McCauley M, et al.; HPTN 052 Study Team. Prevention of HIV-1 infection with early antiretroviral therapy. N Engl J Med. 2011;365(6):493–505. PMID 21767103.",
 "Cohen MS, Chen YQ, McCauley M, et al.; HPTN 052 Study Team. Antiretroviral therapy for the prevention of HIV-1 transmission. N Engl J Med. 2016;375(9):830–839. PMID 27424812.",
 "Quinn TC, Wawer MJ, Sewankambo N, et al. Viral load and heterosexual transmission of HIV type 1. N Engl J Med. 2000;342(13):921–929. PMID 10738050.",
 "Donnell D, Baeten JM, Kiarie J, et al. Heterosexual HIV-1 transmission after initiation of antiretroviral therapy. Lancet. 2010;375(9731):2092–2098. PMID 20537376.",
 "Dellar RC, Dlamini S, Abdool Karim Q. Adolescent girls and young women: key populations for HIV epidemic control. J Int AIDS Soc. 2015;18(2 Suppl 1):19408. PMID 25724504.",
 "Grant RM, Lama JR, Anderson PL, et al.; iPrEx Study Team. Preexposure chemoprophylaxis for HIV prevention in men who have sex with men. N Engl J Med. 2010;363(27):2587–2599. PMID 21091279.",
 "Baeten JM, Donnell D, Ndase P, et al.; Partners PrEP Study Team. Antiretroviral prophylaxis for HIV prevention in heterosexual men and women. N Engl J Med. 2012;367(5):399–410. PMID 22784037.",
 "Delany-Moretlwe S, Hughes JP, Bock P, et al.; HPTN 084 study group. Cabotegravir for the prevention of HIV-1 in women: results from HPTN 084. Lancet. 2022;399(10337):1779–1789. PMID 35378077.",
 "[Clover/Aimee field study] — authors’ own field-study report (cohort, engagement, HIV-testing and PrEP-uptake data, partnership degree, and effect sizes; analytic N = 9,310). [Complete citation.]",
 "Morris M, Kretzschmar M. Concurrent partnerships and the spread of HIV. AIDS. 1997;11(5):641–648. PMID 9108946.",
 "Goodreau SM, Cassels S, Kasprzyk D, et al. Concurrent partnerships, acute infection and HIV epidemic dynamics among young adults in Zimbabwe. AIDS Behav. 2012;16(2):312–322. PMID 21190074.",
 "Maughan-Brown B, Kenyon C, Lurie MN. Partner age differences and concurrency in South Africa: implications for HIV-infection risk among young women. AIDS Behav. 2014;18(12):2469–2476. PMID 25047687.",
 "Stoner MCD, Nguyen N, Kilburn K, et al. Age-disparate partnerships and incident HIV infection in adolescent girls and young women in rural South Africa (HPTN 068). AIDS. 2019;33(1):83–91. PMID 30289813.",
 "Johnson LF. Access to antiretroviral treatment in South Africa, 2004–2011. South Afr J HIV Med. 2012;13(1):22–27.",
 "Celum C, Hosek S, Tsholwana M, et al. PrEP uptake, persistence, adherence among young women in southern Africa (HPTN 082). PLoS Med. 2021;18(6):e1003670. PMID 34143779.",
 "Johnson LF, Dorrington RE. Thembisa version 5.0: a model for evaluating the impact of HIV/AIDS in South Africa. University of Cape Town; 2024. thembisa.org.",
 "Human Sciences Research Council. The Sixth South African National HIV Prevalence, Incidence, Behaviour and Communication Survey (SABSSM VI), 2022. HSRC; 2023.",
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    p.add_run(f"{i}. {r}").font.size = Pt(10)

out = "Aimee_AGYW_HIV_supplement.docx"
doc.save(out)
print("Wrote", out, "with", len(refs), "self-contained references")
